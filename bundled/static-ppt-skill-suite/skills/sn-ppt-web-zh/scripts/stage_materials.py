#!/usr/bin/env python3
"""stage_materials.py —— skill 自带的「附件编排入口」:按 manifest/--input 接收附件 → 拷进 _raw/ →
分类型解析(原生文本/PDF/Office/ODF/图片/媒体/压缩包/未知格式)→ 全文分块与
可视派生物 → 写 catalog.json。material 子代理一条命令跑完常见格式，并对不完整格式获得可执行的建议动作。

用法:
    python stage_materials.py <materials_dir> [attachments.json] [--input PATH ...]

  <materials_dir>   目标目录(通常就是工作区的 `materials/`);附件原件可预先放入
                    `<materials_dir>/_raw/`,或由本脚本按 attachments.json 拷入。
  attachments.json  附件清单(默认读 `<materials_dir>/attachments.json`);每项:
                    {"path": <源路径>, "name": <可选文件名>}  或  直接字符串路径。
                    若不存在,则改为扫描 `<materials_dir>/_raw/` 下已有文件当附件。
  --input PATH      只处理指定附件；可重复。用于多个 Material agent 各自处理互不重叠
                    的附件分片，建议每个 agent 使用独立 `<materials_dir>`。

跨解释器(本脚本用哪个 python 跑都行,内部 worker 会由对应解释器重新调用本文件):
  - 文本解析:`$NORMALIZE_PY stage_materials.py _parse-to-files <file> <out_dir> <name>`
  - 扫描PDF光栅化:`$RASTERIZE_PY stage_materials.py _rasterize <pdf> <out_dir>`

产出:`<materials_dir>/catalog.json` + 各 `<name>.md` + `_chunks/<name>/chunk_NNN.md`
+ `_raw/<name>_pages/pNNN.png`。长文档的 `<name>.md` 保留完整解析文本，chunk 文件提供
确定性、连续且可验收的读取单元，不再截掉尾部。
catalog schema 保持稳定(material.md 依赖这些字段):
    name / raw / ext / kind(doc|image) / text / text_chunks / chars / status /
    coverage / coverage_id / rasterized_pages / note / visual_coverage /
    semantic_coverage / suggested_actions
    所有从附件派生的页图、嵌入图和代表帧都带:
    {is_derivative:true, derived_from:"<原附件名>", derivative_kind:"..."}。
    扫描 PDF 页图另保留 from_scanned_pdf 以表达语义覆盖方式。
末行输出一行 JSON 汇总:{catalog: <path>, entries: N, ok: n_ok, failed: n_failed}。
"""
import sys
import os
import json
import hashlib
import shutil
import subprocess
import logging
import signal
import re
import tempfile
import zipfile
import xml.etree.ElementTree as ET

_IMG_EXT = {"jpg", "jpeg", "png", "webp", "gif", "bmp", "tif", "tiff", "heic", "heif", "svg"}
_SELF = os.path.abspath(__file__)
_OOXML_EXT = {"docx", "pptx", "xlsx", "xlsm"}
_ODF_EXT = {"odt", "ods", "odp"}
_LEGACY_OFFICE_EXT = {"doc", "ppt", "xls", "rtf"}
_CONVERTED_DOC_EXT = _OOXML_EXT | _ODF_EXT | _LEGACY_OFFICE_EXT
_PLAIN_TEXT_EXT = {
    "md", "markdown", "txt", "rst", "csv", "tsv", "json", "jsonl", "ndjson",
    "yaml", "yml", "xml", "html", "htm", "log", "ini", "cfg", "toml",
    "py", "js", "jsx", "ts", "tsx", "css", "sql",
}
_AUDIO_EXT = {"mp3", "wav", "m4a", "aac", "flac", "ogg", "opus"}
_VIDEO_EXT = {"mp4", "mov", "mkv", "webm", "avi", "m4v", "mpeg", "mpg"}
_ARCHIVE_EXT = {"zip"}
_CHUNK_CHARS = max(1000, int(os.environ.get("MATERIALS_CHUNK_CHARS", "12000")))
_PDF_TIMEOUT = 90
_MARKITDOWN = None
_MARKITDOWN_UNAVAILABLE = False

# 解析用解释器(第三方库不在 skill 树,靠外部 venv;见 install.sh / SKILL 环境依赖段)
NORMALIZE_PY = os.environ.get("NORMALIZE_PY", sys.executable)
# 光栅化用解释器(需 PyMuPDF;默认回退当前解释器,装了 pymupdf 就能用)
RASTERIZE_PY = os.environ.get("RASTERIZE_PY", os.environ.get("PYMUPDF_PY", sys.executable))


def _normalize_worker_init():
    global _MARKITDOWN, _MARKITDOWN_UNAVAILABLE
    if _MARKITDOWN_UNAVAILABLE:
        return False
    for name in ("pdfminer", "pdfminer.pdffont", "pdfminer.pdfinterp", "pdfminer.layout"):
        logging.getLogger(name).setLevel(logging.ERROR)
    try:
        from markitdown import MarkItDown
        _MARKITDOWN = MarkItDown()
        return True
    except Exception:
        _MARKITDOWN_UNAVAILABLE = True
        return False


def _docx_fallback(path):
    try:
        from docx import Document
        document = Document(path)
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        for section in document.sections:
            for container in (section.header, section.footer):
                lines.extend(p.text for p in container.paragraphs if p.text.strip())
        text = "\n".join(lines).strip()
        if text:
            return text
    except Exception:
        pass
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        with zipfile.ZipFile(path) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))
        return "\n".join(
            "".join(node.text for node in paragraph.iter(namespace + "t") if node.text)
            for paragraph in root.iter(namespace + "p")
        ).strip()
    except Exception:
        return ""


def _xlsx_fallback(path):
    try:
        import openpyxl
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"## {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell) for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        workbook.close()
        return "\n".join(lines).strip()
    except Exception:
        return ""


def _pptx_fallback(path):
    """Extract slide and notes text from OOXML without python-pptx/MarkItDown."""
    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            slide_names = sorted(
                (name for name in names if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)),
                key=lambda name: int(re.search(r"(\d+)", os.path.basename(name)).group(1)),
            )
            note_names = sorted(
                (name for name in names if re.fullmatch(r"ppt/notesSlides/notesSlide\d+\.xml", name)),
                key=lambda name: int(re.search(r"(\d+)", os.path.basename(name)).group(1)),
            )
            blocks = []
            for index, name in enumerate(slide_names, 1):
                root = ET.fromstring(archive.read(name))
                texts = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
                blocks.append(f"[Slide {index}]\n" + "\n".join(texts))
            for index, name in enumerate(note_names, 1):
                root = ET.fromstring(archive.read(name))
                texts = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
                if texts:
                    blocks.append(f"[Notes {index}]\n" + "\n".join(texts))
            return "\n\n".join(blocks).strip()
    except Exception:
        return ""


def _odf_fallback(path):
    try:
        with zipfile.ZipFile(path) as archive:
            root = ET.fromstring(archive.read("content.xml"))
        lines = []
        for node in root.iter():
            if node.text and node.text.strip():
                lines.append(node.text.strip())
        return "\n".join(lines).strip()
    except Exception:
        return ""


def _rtf_fallback(path):
    """Best-effort readable text for simple RTF; complex RTF still prefers LibreOffice."""
    try:
        raw = _plain_text(path)
        raw = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: bytes.fromhex(m.group(1)).decode("cp1252", "replace"), raw)
        raw = re.sub(r"\\(?:par|line)\b", "\n", raw)
        raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
        raw = raw.replace("{", "").replace("}", "")
        return raw.strip()
    except Exception:
        return ""


def _markitdown(path):
    global _MARKITDOWN
    available = _MARKITDOWN is not None or _normalize_worker_init()
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    fallback = {
        "docx": _docx_fallback,
        "xlsx": _xlsx_fallback,
        "xlsm": _xlsx_fallback,
        "pptx": _pptx_fallback,
        "odt": _odf_fallback,
        "ods": _odf_fallback,
        "odp": _odf_fallback,
        "rtf": _rtf_fallback,
    }.get(ext)
    if available:
        try:
            text = _MARKITDOWN.convert(path).text_content or ""
            if text.strip():
                return text
        except Exception:
            pass
    return fallback(path) if fallback else ""


def _pdf_text(path):
    def timeout(*_args):
        raise TimeoutError()

    signal.signal(signal.SIGALRM, timeout)
    signal.alarm(_PDF_TIMEOUT)
    try:
        try:
            from pdfminer.high_level import extract_text
            return extract_text(path) or ""
        except Exception:
            from pypdf import PdfReader
            return "\n\n".join((page.extract_text() or "") for page in PdfReader(path).pages)
    finally:
        signal.alarm(0)


def _finish_text(text):
    text = (text or "").strip()
    size = len(text)
    if not size:
        return {"status": "failed", "content_chars": 0,
                "content": "", "note": "empty parse"}
    return {"status": "ok", "content_chars": size, "content": text}


def _plain_text(path):
    """Read text-native attachments without pulling in document converters."""
    last_error = None
    for encoding in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            with open(path, encoding=encoding) as stream:
                return stream.read()
        except (UnicodeError, OSError) as exc:
            last_error = exc
    raise last_error or UnicodeError("unable to decode plain-text attachment")


def _safe_name(name):
    base = os.path.basename(name or "").strip()
    base = re.sub(r"[^A-Za-z0-9._()\- \u4e00-\u9fff]+", "_", base)
    return base[:180] or "attachment"


def _office_binary():
    return shutil.which("libreoffice") or shutil.which("soffice")


def _embedded_media(path, out_dir, source_name):
    """Extract OOXML/ODF embedded raster images as separately reviewable assets."""
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    prefixes = {
        "docx": ("word/media/",),
        "pptx": ("ppt/media/",),
        "xlsx": ("xl/media/",),
        "xlsm": ("xl/media/",),
        "odt": ("Pictures/",),
        "ods": ("Pictures/",),
        "odp": ("Pictures/",),
    }.get(ext, ())
    if not prefixes:
        return []
    outputs = []
    try:
        with zipfile.ZipFile(path) as archive:
            members = [
                name for name in archive.namelist()
                if any(name.startswith(prefix) for prefix in prefixes)
                and os.path.splitext(name)[1].lstrip(".").lower() in _IMG_EXT - {"svg", "heic", "heif"}
            ]
            os.makedirs(out_dir, exist_ok=True)
            for index, member in enumerate(sorted(members), 1):
                suffix = os.path.splitext(member)[1].lower() or ".png"
                target = os.path.join(out_dir, f"media_{index:03d}{suffix}")
                with open(target, "wb") as stream:
                    stream.write(archive.read(member))
                outputs.append({
                    "name": f"{source_name} · embedded {index}",
                    "raw": _workspace_path(target),
                    "ext": suffix.lstrip("."),
                    "kind": "image",
                    "status": "ok",
                    "from_document": source_name,
                    "is_derivative": True,
                    "derived_from": source_name,
                    "derivative_kind": "embedded_media",
                    "coverage": {"status": "complete", "unit": "asset", "covered": 1, "total": 1},
                })
    except Exception:
        return []
    return outputs


def _render_office_pages(path, out_dir):
    """Use an existing LibreOffice installation; never install from an agent turn."""
    office = _office_binary()
    if not office:
        return [], 0, "LibreOffice/soffice unavailable"
    try:
        with tempfile.TemporaryDirectory() as temporary:
            result = subprocess.run(
                [office, "--headless", "--convert-to", "pdf", "--outdir", temporary, path],
                capture_output=True, text=True, timeout=120,
            )
            pdfs = sorted(
                os.path.join(temporary, item) for item in os.listdir(temporary)
                if item.lower().endswith(".pdf")
            )
            if result.returncode != 0 or not pdfs:
                detail = (result.stderr or result.stdout or "conversion produced no PDF").strip()
                return [], 0, detail[:240]
            pages, total = _rasterize(pdfs[0], out_dir)
            return pages, total, ""
    except Exception as exc:
        return [], 0, f"{type(exc).__name__}: {exc}"


def _normalize_visual_attachment(path, out_dir):
    """Normalize formats that vision backends do not reliably accept into PNG frames."""
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    if ext in {"jpg", "jpeg", "png", "webp"}:
        return [path], ""
    os.makedirs(out_dir, exist_ok=True)
    if ext == "svg":
        try:
            import cairosvg
            target = os.path.join(out_dir, "frame_001.png")
            cairosvg.svg2png(url=path, write_to=target)
            return [target], ""
        except Exception as exc:
            return [], f"SVG rasterizer unavailable: {type(exc).__name__}: {exc}"
    try:
        from PIL import Image, ImageSequence
        image = Image.open(path)
        outputs = []
        for index, frame in enumerate(ImageSequence.Iterator(image), 1):
            if index > int(os.environ.get("MATERIAL_IMAGE_MAX_FRAMES", "40")):
                break
            target = os.path.join(out_dir, f"frame_{index:03d}.png")
            frame.convert("RGBA").save(target)
            outputs.append(target)
        return outputs, "" if outputs else "image contained no readable frames"
    except Exception as exc:
        return [], f"image normalization failed: {type(exc).__name__}: {exc}"


def _media_probe(path):
    probe = shutil.which("ffprobe")
    if not probe:
        return "", "ffprobe unavailable"
    try:
        result = subprocess.run(
            [probe, "-v", "quiet", "-print_format", "json", "-show_format", "-show_streams", path],
            capture_output=True, text=True, timeout=60,
        )
        return result.stdout.strip(), "" if result.returncode == 0 else (result.stderr or "ffprobe failed")[:200]
    except Exception as exc:
        return "", f"{type(exc).__name__}: {exc}"


def _video_frames(path, out_dir):
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        return [], "ffmpeg unavailable"
    os.makedirs(out_dir, exist_ok=True)
    target = os.path.join(out_dir, "frame_%03d.png")
    try:
        result = subprocess.run(
            [ffmpeg, "-nostdin", "-i", path, "-vf", "fps=1/30,scale=1280:-2", "-frames:v", "12", target],
            capture_output=True, text=True, timeout=180,
        )
        frames = sorted(os.path.join(out_dir, name) for name in os.listdir(out_dir) if name.endswith(".png"))
        return frames, "" if frames else (result.stderr or "no video frames extracted")[-300:]
    except Exception as exc:
        return [], f"{type(exc).__name__}: {exc}"


def _sha256_text(text):
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _write_text_outputs(text, out_dir, name):
    """Write a lossless full-text file plus contiguous, hash-addressed chunks."""
    text = text or ""
    os.makedirs(out_dir, exist_ok=True)
    safe_name = os.path.basename(name) or "material"
    full_path = os.path.join(out_dir, safe_name + ".md")
    with open(full_path, "w", encoding="utf-8") as stream:
        stream.write(text)

    chunk_dir = os.path.join(out_dir, "_chunks", safe_name)
    os.makedirs(chunk_dir, exist_ok=True)
    for old_name in os.listdir(chunk_dir):
        if old_name.startswith("chunk_") and old_name.endswith(".md"):
            os.unlink(os.path.join(chunk_dir, old_name))

    chunks = []
    for number, start in enumerate(range(0, len(text), _CHUNK_CHARS), 1):
        end = min(len(text), start + _CHUNK_CHARS)
        content = text[start:end]
        path = os.path.join(chunk_dir, f"chunk_{number:03d}.md")
        with open(path, "w", encoding="utf-8") as stream:
            stream.write(content)
        chunks.append({
            "path": path,
            "start_char": start,
            "end_char": end,
            "chars": len(content),
            "sha256": _sha256_text(content),
        })
    return {
        "text_path": full_path,
        "text_sha256": _sha256_text(text),
        "text_chunks": chunks,
    }


def _parse_to_files(path, out_dir, name):
    record = _parse_one(path)
    content = record.pop("content", "")
    if record.get("status") == "ok" and content:
        record.update(_write_text_outputs(content, out_dir, name))
    return record


def _parse_one(path):
    ext = os.path.splitext(path)[1].lstrip(".").lower()
    if ext in _PLAIN_TEXT_EXT:
        try:
            return _finish_text(_plain_text(path))
        except Exception as exc:
            return {"status": "failed", "content": "", "note": f"plain text: {exc}"}
    if ext == "pdf":
        record = _finish_text(_pdf_text(path))
        if record.get("content_chars", 0) < 20:
            record.update(status="failed", note="near-empty (likely scanned PDF, needs vision)")
        return record
    if ext in _CONVERTED_DOC_EXT:
        text = _markitdown(path)
        if text.strip():
            return _finish_text(text)
        return {
            "status": "failed", "content": "",
            "note": f"{ext} text extraction unavailable; existing LibreOffice/MarkItDown may be used",
            "suggested_actions": ["libreoffice_to_pdf", "markitdown", "manual_visual_review"],
        }
    if ext in _AUDIO_EXT | _VIDEO_EXT:
        metadata, note = _media_probe(path)
        if metadata:
            record = _finish_text(metadata)
            record["note"] = "media metadata only; spoken content still requires transcription"
            record["semantic_coverage"] = "incomplete"
            return record
        return {
            "status": "failed", "content": "", "note": note,
            "suggested_actions": ["existing_ffprobe_ffmpeg", "existing_asr", "manual_review"],
        }
    if ext in _ARCHIVE_EXT:
        try:
            with zipfile.ZipFile(path) as archive:
                names = [name for name in archive.namelist() if not name.endswith("/")]
            record = _finish_text("[Archive members]\n" + "\n".join(names))
            record.update(
                semantic_coverage="incomplete",
                note="archive inventory only; extract safe members and stage each relevant file",
                suggested_actions=["safe_extract_in_assignment_dir", "stage_each_member"],
            )
            return record
        except Exception as exc:
            return {"status": "failed", "content": "", "note": f"archive: {exc}"}
    return {
        "status": "unsupported", "content": "", "note": f"unsupported attachment type ({ext or 'no extension'})",
        "suggested_actions": ["inspect_file_signature", "use_existing_converter", "derive_reviewable_text_or_images"],
    }


def _rasterize_worker(pdf_path, out_dir):
    # 0 means all pages. A positive operational cap is allowed, but the catalog
    # will mark the source incomplete and deterministic acceptance must reject it.
    max_pages = int(os.environ.get("MATERIAL_RASTER_MAX_PAGES", "0"))
    dpi = int(os.environ.get("MATERIAL_RASTER_DPI", "150"))
    max_pixels = int(os.environ.get("MATERIAL_RASTER_MAXPX", "2600"))
    try:
        import fitz
    except Exception as exc:
        return {"pages": [], "total": 0, "note": f"PyMuPDF unavailable: {exc}"}
    pages = []
    total = 0
    try:
        document = fitz.open(pdf_path)
        total = document.page_count
        os.makedirs(out_dir, exist_ok=True)
        page_limit = total if max_pages <= 0 else min(total, max_pages)
        for index in range(page_limit):
            page = document[index]
            zoom = dpi / 72.0
            longest = max(page.rect.width, page.rect.height) or 1
            if longest * zoom > max_pixels:
                zoom = max_pixels / longest
            target = os.path.join(out_dir, f"p{index + 1:03d}.png")
            page.get_pixmap(matrix=fitz.Matrix(zoom, zoom)).save(target)
            pages.append(target)
        document.close()
    except Exception as exc:
        return {"pages": pages, "total": total, "note": f"{type(exc).__name__}: {str(exc)[:150]}"}
    return {"pages": pages, "total": total}


def _load_attachments(mdir):
    """确定附件清单:优先 attachments.json;否则扫 _raw/ 下已有文件。"""
    aj = os.path.join(mdir, "attachments.json")
    if os.path.exists(aj):
        try:
            data = json.load(open(aj, encoding="utf-8"))
            if isinstance(data, dict):
                data = data.get("attachments") or data.get("files") or []
            return data or []
        except Exception:
            pass
    raw = os.path.join(mdir, "_raw")
    if os.path.isdir(raw):
        return [{"path": os.path.join(raw, n), "name": n}
                for n in sorted(os.listdir(raw))
                if os.path.isfile(os.path.join(raw, n)) and not n.endswith("_pages")]
    return []


def stage(mdir, attachments):
    raw = os.path.join(mdir, "_raw")
    os.makedirs(raw, exist_ok=True)
    catalog = []
    for a in attachments:
        # attachments.json 的路径字段兼容 raw 与 path——都接,
        # 别只认 path(否则 raw 项 → src=None → 误判 missing、白丢已拷进 _raw 的原件)。
        src = (a.get("raw") or a.get("path")) if isinstance(a, dict) else a
        name = _safe_name((a.get("name") if isinstance(a, dict) else None)
                          or (os.path.basename(src) if src else "unknown"))
        if not src or not os.path.exists(src):
            catalog.append({"name": name, "status": "missing"})
            continue
        dst = os.path.join(raw, name)
        try:
            if os.path.abspath(src) != os.path.abspath(dst):
                shutil.copy2(src, dst)
        except Exception as e:
            catalog.append({"name": name, "status": "failed", "note": f"copy: {e}"})
            continue
        ext = os.path.splitext(name)[1].lstrip(".").lower()
        source_sha256 = _sha256_file(dst)
        entry = {"name": name, "raw": _workspace_path(dst), "ext": ext,
                 "source_sha256": source_sha256}
        derivative_entries = []
        if ext in _IMG_EXT:
            frames, note = _normalize_visual_attachment(dst, os.path.join(raw, name + "_frames"))
            if frames:
                entry.update(
                    kind="image",
                    status="ok",
                    source_raw=_workspace_path(dst),
                    raw=_workspace_path(frames[0]),
                    coverage={"status": "complete", "unit": "frames", "covered": len(frames), "total": len(frames)},
                    normalized_frames=[_workspace_path(frame) for frame in frames],
                )
                derivative_entries.extend({
                    "name": f"{name} · frame {index}",
                    "raw": _workspace_path(frame),
                    "ext": "png", "kind": "image", "status": "ok",
                    "frame": index, "from_image": name,
                    "is_derivative": True, "derived_from": name,
                    "derivative_kind": "normalized_frame",
                    "coverage": {"status": "complete", "unit": "asset", "covered": 1, "total": 1},
                } for index, frame in enumerate(frames[1:], 2))
            else:
                entry.update(
                    kind="image", status="incomplete", note=note,
                    coverage={"status": "incomplete", "unit": "frames", "covered": 0, "total": 1},
                    suggested_actions=["use_existing_image_converter", "browser_rasterize", "manual_visual_review"],
                )
        else:
            try:
                out = subprocess.run([NORMALIZE_PY, _SELF, "_parse-to-files", dst, mdir, name],
                                     capture_output=True, text=True, timeout=180)
                rec = json.loads(out.stdout.strip().splitlines()[-1])
                if rec.get("text_path"):
                    chunks = []
                    for chunk in rec.get("text_chunks") or []:
                        item = dict(chunk)
                        item["path"] = _workspace_path(item["path"])
                        chunks.append(item)
                    chars = int(rec.get("content_chars") or 0)
                    entry.update(
                        kind="doc",
                        text=_workspace_path(rec["text_path"]),
                        text_sha256=rec.get("text_sha256"),
                        text_chunks=chunks,
                        chars=chars,
                        status="ok",
                        coverage={
                            "status": "complete",
                            "unit": "chars",
                            "covered": sum(int(item.get("chars") or 0) for item in chunks),
                            "total": chars,
                            "chunks": len(chunks),
                        },
                    )
                    for key in ("note", "semantic_coverage", "suggested_actions"):
                        if rec.get(key) is not None:
                            entry[key] = rec[key]
                    if rec.get("semantic_coverage") == "incomplete":
                        entry["status"] = "incomplete"
                        entry["coverage"]["status"] = "incomplete"
                else:
                    entry.update(
                        kind="doc", status=rec.get("status", "failed"), note=rec.get("note"),
                        suggested_actions=rec.get("suggested_actions") or [],
                    )
                    # 图片式/扫描 PDF 抽不出文本 → 光栅化页图,当图片交 vision_analyze(否则内容彻底丢失)
                    if ext == "pdf":
                        _pgs, _total = _rasterize(dst, os.path.join(raw, name + "_pages"))
                        if _pgs:
                            complete = bool(_total) and len(_pgs) == _total
                            entry["status"] = "ok" if complete else "incomplete"
                            entry["coverage"] = {
                                "status": "complete" if complete else "incomplete",
                                "unit": "pages",
                                "covered": len(_pgs),
                                "total": _total,
                            }
                            entry["rasterized_pages"] = len(_pgs)
                            entry["note"] = ("图片式/扫描 PDF、无文本层;已 rasterize "
                                             + str(len(_pgs)) + (f"/{_total}" if _total else "")
                                             + " 页成图,内容见下方同名 image 条目,请用 vision_analyze 逐页读。")
                            derivative_entries = [{"name": f"{name} · p{i}",
                                                   "raw": _workspace_path(pg),
                                                   "ext": "png", "kind": "image", "status": "ok",
                                                   "page": i, "from_scanned_pdf": name,
                                                   "is_derivative": True, "derived_from": name,
                                                   "derivative_kind": "scanned_pdf_page",
                                                   "material_asset_type": "page_context",
                                                   "reuse_policy": "reference_only_until_cropped",
                                                   "coverage": {"status": "complete", "unit": "asset", "covered": 1, "total": 1}}
                                                  for i, pg in enumerate(_pgs, 1)]
            except Exception as e:
                entry.update(kind="doc", status="failed", note=f"{type(e).__name__}: {e}")

            # A text layer does not preserve charts, pictures, page composition or slide/sheet layout.
            # Keep those visual facts as reviewable page images whenever existing converters allow it.
            if ext == "pdf" and entry.get("status") == "ok" and not entry.get("rasterized_pages"):
                pages, total = _rasterize(dst, os.path.join(raw, name + "_pages"))
                if pages:
                    entry["visual_coverage"] = {
                        "status": "complete" if total and len(pages) == total else "incomplete",
                        "covered": len(pages), "total": total, "unit": "pages",
                    }
                    derivative_entries.extend({
                        "name": f"{name} · visual p{index}",
                        "raw": _workspace_path(page),
                        "ext": "png", "kind": "image", "status": "ok",
                        "page": index, "from_pdf": name,
                        "is_derivative": True, "derived_from": name,
                        "derivative_kind": "pdf_page_visual",
                        "material_asset_type": "page_context",
                        "reuse_policy": "reference_only_until_cropped",
                        "coverage": {"status": "complete", "unit": "asset", "covered": 1, "total": 1},
                    } for index, page in enumerate(pages, 1))
                else:
                    entry["visual_coverage"] = {"status": "unavailable", "covered": 0, "total": total, "unit": "pages"}
                    entry["note"] = ((entry.get("note") or "") + "; PDF page rasterization unavailable").strip("; ")

            if ext in _OOXML_EXT | _ODF_EXT:
                derivative_entries.extend(_embedded_media(
                    dst, os.path.join(raw, name + "_media"), name
                ))
                pages, total, note = _render_office_pages(dst, os.path.join(raw, name + "_pages"))
                if pages:
                    entry["visual_coverage"] = {
                        "status": "complete" if total and len(pages) == total else "incomplete",
                        "covered": len(pages), "total": total, "unit": "pages",
                    }
                    derivative_entries.extend({
                        "name": f"{name} · rendered p{index}",
                        "raw": _workspace_path(page), "ext": "png", "kind": "image", "status": "ok",
                        "page": index, "from_document": name,
                        "is_derivative": True, "derived_from": name,
                        "derivative_kind": "document_page_visual",
                        "coverage": {"status": "complete", "unit": "asset", "covered": 1, "total": 1},
                    } for index, page in enumerate(pages, 1))
                elif note:
                    entry["visual_coverage"] = {"status": "unavailable", "covered": 0, "unit": "pages"}
                    entry["note"] = ((entry.get("note") or "") + "; page rendering: " + note).strip("; ")

            if ext in _LEGACY_OFFICE_EXT and entry.get("status") != "ok":
                pages, total, note = _render_office_pages(dst, os.path.join(raw, name + "_pages"))
                if pages:
                    entry.update(
                        status="ok",
                        coverage={"status": "complete" if total and len(pages) == total else "incomplete",
                                  "unit": "pages", "covered": len(pages), "total": total},
                        visual_coverage={"status": "complete" if total and len(pages) == total else "incomplete",
                                         "unit": "pages", "covered": len(pages), "total": total},
                        note="legacy office document rendered to page images; use vision_analyze",
                    )
                    derivative_entries.extend({
                        "name": f"{name} · rendered p{index}",
                        "raw": _workspace_path(page), "ext": "png", "kind": "image", "status": "ok",
                        "page": index, "from_document": name,
                        "is_derivative": True, "derived_from": name,
                        "derivative_kind": "document_page_visual",
                        "coverage": {"status": "complete", "unit": "asset", "covered": 1, "total": 1},
                    } for index, page in enumerate(pages, 1))
                else:
                    entry["note"] = ((entry.get("note") or "") + "; page rendering: " + note).strip("; ")

            if ext in _VIDEO_EXT:
                frames, note = _video_frames(dst, os.path.join(raw, name + "_frames"))
                derivative_entries.extend({
                    "name": f"{name} · sampled frame {index}",
                    "raw": _workspace_path(frame), "ext": "png", "kind": "image", "status": "ok",
                    "frame": index, "from_video": name,
                    "is_derivative": True, "derived_from": name,
                    "derivative_kind": "sampled_video_frame",
                    "coverage": {"status": "sampled", "unit": "frame", "covered": 1, "total": None},
                } for index, frame in enumerate(frames, 1))
                if note:
                    entry["note"] = ((entry.get("note") or "") + "; frame extraction: " + note).strip("; ")
        entry["coverage_id"] = _coverage_id(entry)
        catalog.append(entry)
        for derivative in derivative_entries:
            derivative["coverage_id"] = _coverage_id(derivative)
        catalog.extend(derivative_entries)
    with open(os.path.join(mdir, "catalog.json"), "w", encoding="utf-8") as f:
        json.dump(catalog, f, ensure_ascii=False, indent=2)
    return catalog


def _workspace_path(path):
    """Return a stable workspace-relative path for catalogs in nested shards."""
    return os.path.relpath(os.path.abspath(path), os.getcwd()).replace(os.sep, "/")


def _sha256_file(path):
    digest = hashlib.sha256()
    with open(path, "rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _coverage_id(entry):
    payload = {
        "name": entry.get("name"),
        "source_sha256": entry.get("source_sha256"),
        "status": entry.get("status"),
        "coverage": entry.get("coverage"),
        "text_sha256": entry.get("text_sha256"),
        "chunk_sha256": [item.get("sha256") for item in entry.get("text_chunks") or []],
    }
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()[:20]


def _rasterize(pdf_path, out_dir):
    """Use the PyMuPDF interpreter to run this file's isolated raster worker."""
    try:
        out = subprocess.run([RASTERIZE_PY, _SELF, "_rasterize", pdf_path, out_dir],
                             capture_output=True, text=True, timeout=300)
        rec = json.loads(out.stdout.strip().splitlines()[-1])
        return rec.get("pages") or [], rec.get("total") or 0
    except Exception:
        return [], 0


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "_parse-one":
        try:
            payload = _parse_one(sys.argv[2])
        except Exception as exc:
            payload = {"status": "failed", "content": "",
                       "note": f"{type(exc).__name__}: {str(exc)[:150]}"}
        print(json.dumps(payload, ensure_ascii=False))
        sys.exit(0)
    if len(sys.argv) >= 5 and sys.argv[1] == "_parse-to-files":
        try:
            payload = _parse_to_files(sys.argv[2], sys.argv[3], sys.argv[4])
        except Exception as exc:
            payload = {"status": "failed", "content_chars": 0,
                       "note": f"{type(exc).__name__}: {str(exc)[:150]}"}
        print(json.dumps(payload, ensure_ascii=False))
        sys.exit(0)
    if len(sys.argv) >= 4 and sys.argv[1] == "_rasterize":
        try:
            payload = _rasterize_worker(sys.argv[2], sys.argv[3])
        except Exception as exc:
            payload = {"pages": [], "total": 0,
                       "note": f"{type(exc).__name__}: {str(exc)[:150]}"}
        print(json.dumps(payload, ensure_ascii=False))
        sys.exit(0)
    if len(sys.argv) == 2 and sys.argv[1] in {"-h", "--help"}:
        print("usage: stage_materials.py MATERIALS_DIR [attachments.json] [--input PATH ...]")
        sys.exit(0)
    if len(sys.argv) < 2:
        print(json.dumps({"error": "usage: stage_materials.py <materials_dir> [attachments.json]"}))
        sys.exit(0)
    mdir = sys.argv[1]
    os.makedirs(mdir, exist_ok=True)
    args = sys.argv[2:]
    explicit_inputs = []
    positional = []
    i = 0
    while i < len(args):
        if args[i] == "--input" and i + 1 < len(args):
            explicit_inputs.append(args[i + 1])
            i += 2
            continue
        positional.append(args[i])
        i += 1
    atts = ([{"path": path, "name": os.path.basename(path)} for path in explicit_inputs]
            if explicit_inputs else _load_attachments(mdir))
    if not explicit_inputs and positional and os.path.exists(positional[0]):
        try:
            data = json.load(open(positional[0], encoding="utf-8"))
            atts = (data.get("attachments") if isinstance(data, dict) else data) or atts
        except Exception:
            pass
    cat = stage(mdir, atts)
    n_ok = sum(1 for e in cat if e.get("status") == "ok")
    n_failed = sum(1 for e in cat if e.get("status") in ("failed", "missing", "incomplete", "unsupported"))
    print(json.dumps({"catalog": os.path.join(mdir, "catalog.json"),
                      "entries": len(cat), "ok": n_ok, "failed": n_failed}, ensure_ascii=False))
